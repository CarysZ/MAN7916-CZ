import os
import re
import matplotlib.pyplot as plt
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import tomotopy as tp
from docx import Document
from wordcloud import WordCloud

def preprocess(text: str) -> list:
    stop_words = set(stopwords.words('english'))
    text = re.sub(r'[^A-Za-z]', ' ', text)
    words = word_tokenize(text.lower())
    words = [word for word in words if word not in stop_words and word.isalpha()]
    return words

def load_corpus(directory: str) -> list:
    texts = []
    for filename in os.listdir(directory):
        if filename.endswith('.txt'):
            with open(os.path.join(directory, filename), 'r', encoding='utf-8') as f:
                texts.append(f.read())
    return texts

def train_lda_model(preprocessed: list, num_topics: int, alpha=0.1, beta=0.01) -> tuple:
    mdl = tp.LDAModel(
        tw=tp.TermWeight.IDF,
        k=num_topics,
        alpha=alpha,
        eta=beta,
        rm_top=10,
        min_cf=2,
        min_df=1
    )
    for doc in preprocessed:
        if doc:
            mdl.add_doc(doc)

    mdl.train(1000)

    # Compute u_mass coherence score
    coherence = tp.coherence.Coherence(mdl, coherence='u_mass')
    score = coherence.get_score()
    return mdl, score

# Load and preprocess corpus
ELcorpus = load_corpus('local_data/newsgroup')
preprocessed = [preprocess(text) for text in ELcorpus]

topic_ranges = range(2, 11)
coherence_scores = []

# Create Word document
doc = Document()
doc.add_heading('Coherence Scores of Topic Models', 5)

# Train models and store results
for n_topics in topic_ranges:
    print(f"Training model with {n_topics} topics...")
    _, score = train_lda_model(preprocessed, n_topics)
    coherence_scores.append(score)

    doc.add_paragraph(f'Number of topics: {n_topics}, Coherence score: {score:.4f}')

# Save Word document
doc.save('assignments/submissions/assignment_5/topic_model_results.docx')

# Plot u_mass coherence scores
plt.figure(figsize=(12, 8))
plt.plot(list(topic_ranges), coherence_scores, label='Coherence', linewidth=1.5)
plt.xlabel('Number of Topics')
plt.ylabel('Coherence Score')
plt.title('Coherence Scores for Topic Models')
plt.legend(loc='upper left')
plt.grid(False)
plt.savefig('assignments/submissions/assignment_5/coherence_scores.jpg')

# Tinker with alpha and beta to improve the best model
best_num_topics = 10
best_score = -float('inf')
best_model = None

print("Tinkering with hyperparameters...")
for alpha in [0.05, 0.1, 0.2]:
    for beta in [0.01, 0.05, 0.1]:
        print(f"Trying alpha={alpha}, beta={beta}")
        mdl, score = train_lda_model(preprocessed, best_num_topics, alpha, beta)
        
        if score > best_score:
            best_score = score
            best_model = mdl 

doc.add_paragraph(f'Best model coherence score: {best_score:.4f}')
doc.save('assignments/submissions/assignment_5/topic_model_results.docx')

# Create word clouds for the best model
def create_wordclouds(model: tp.LDAModel, num_topics: int):
    """Create word clouds for each topic and save them as images."""
    for k in range(num_topics):
        words = model.get_topic_words(k, top_n=10)
        word_freq = {word: prob for word, prob in words}
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate_from_frequencies(word_freq)
        wordcloud.to_file(f'assignments/submissions/assignment_5/topic_{k}_wordcloud.jpg')

print("Creating word clouds...")
create_wordclouds(best_model, best_num_topics)

#Infer the topics of the new article
unseen_file = 'local_data/ethico.txt'
with open(unseen_file, 'r', encoding='utf-8') as f:
    unseen_article = f.read()

preprocessed_unseen = preprocess(unseen_article)
topics = best_model.infer(best_model.make_doc(preprocessed_unseen))
doc.add_paragraph('Topics for unseen article:')
for topic_id, prob in enumerate(topics[0]):
    doc.add_paragraph(f'Topic {topic_id}: {prob:.4f}')
doc.save('assignments/submissions/assignment_5/topic_model_results.docx')