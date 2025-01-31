import numpy as np
import csv
from pathlib import Path
from gensim.models import KeyedVectors
from scipy.spatial.distance import cosine
from docx import Document
from pprint import pprint
from itertools import combinations

# Paths setup
local_data_path = Path(__file__).resolve().parent.parent.parent.parent / "local_data"
glove_file_path = local_data_path / "glove.6B.100d.txt"
output_docx_path = "assignments/submissions/assignment_5/word_embedding_results.docx"
article_preprint_path = "assignments/materials/week_4/article_preprint.txt"
output_csv_path = "assignments/submissions/assignment_5/word_list_for_evaluation.csv"

# Load GloVe embeddings
print("Loading GloVe model...")
embeddings = KeyedVectors.load_word2vec_format(
    glove_file_path.as_posix(), binary=False, no_header=True)

# Define root words
root_words = ["entrepreneurial", "creative", "innovative", "trailblazing"]

# Display most similar words for each root word
print("Most similar words to each root word:")
for word in root_words:
    if word in embeddings:
        similar_words = embeddings.similar_by_word(word, topn=5)
        print(f"Most similar words to '{word}':")
        pprint(similar_words, indent=4)
    else:
        print(f"'{word}' not found in GloVe embeddings.")

# Compute cosine similarities between root words
document = Document()
document.add_heading("Word Embedding Analysis Results", level=1)

cosine_similarities = []
for idx, (i, j) in enumerate(combinations(range(len(root_words)), 2), 1):
    word1, word2 = root_words[i], root_words[j]
    similarity = 1 - cosine(embeddings[word1], embeddings[word2])
    cosine_similarities.append((word1, word2, similarity))
    print(f"Similarity between '{word1}' and '{word2}': {similarity:.4f}")

# Save cosine similarities to Word document
document.add_heading("Cosine Similarities Between Root Words", level=2)
for word1, word2, similarity in cosine_similarities:
    document.add_paragraph(f"{word1} - {word2}: {similarity:.4f}")

# Drop the least similar root word
least_fitting_word = min(root_words, key=lambda w: sum(
    1 - cosine(embeddings[w], embeddings[other]) for other in root_words if other != w
))
remaining_root_words = [w for w in root_words if w != least_fitting_word]

# Save remaining root words to Word document
document.add_heading("Remaining Root Words", level=2)
document.add_paragraph(", ".join(remaining_root_words))

# Calculate the average vector of remaining root words
average_vector = np.mean([embeddings[word] for word in remaining_root_words], axis=0)
document.add_heading("First Five Dimensions of Average Vector", level=2)
document.add_paragraph(str(average_vector[:5]))

# Cosine similarity between average vector and each root word
document.add_heading("Cosine Similarities with Average Vector", level=2)
for word in remaining_root_words:
    similarity = 1 - cosine(average_vector, embeddings[word])
    document.add_paragraph(f"{word}: {similarity:.4f}")

# Deductive word list (50 most similar words)
document.add_heading("Deductive Word List", level=2)
deductive_word_list = embeddings.similar_by_vector(average_vector, topn=50)
for word, score in deductive_word_list:
    document.add_paragraph(f"{word}: {score:.4f}")

# Read text from article_preprint.txt
with open(article_preprint_path, "r", encoding="utf-8") as file:
    text = file.read().lower().split()

# Identify 50 most similar words from article_preprint.txt
unique_words_in_text = set(text).intersection(set(embeddings.key_to_index.keys()))
inductive_word_list = []
for word in unique_words_in_text:
    similarity = 1 - cosine(average_vector, embeddings[word])
    inductive_word_list.append((word, similarity))

# Sort and select top 50 words
inductive_word_list = sorted(inductive_word_list, key=lambda x: x[1], reverse=True)[:50]

# Save inductive word list to the document
document.add_heading("Inductive Word List", level=2)
for word, score in inductive_word_list:
    document.add_paragraph(f"{word}: {score:.4f}")

# Combine deductive and inductive lists
combined_word_list = {word: score for word, score in deductive_word_list}
combined_word_list.update({word: score for word, score in inductive_word_list})

# Save combined list to CSV
with open(output_csv_path, mode="w", newline="", encoding="utf-8") as file:
    writer = csv.writer(file)
    writer.writerow(["word", "score", "eval"])
    for word, score in combined_word_list.items():
        writer.writerow([word, score, ""])  # eval column is empty

# Save the results to the Word document
document.save(output_docx_path)

print("Saved word embedding results to word_embedding_results.docx")
print("Saved combined word list to word_list_for_evaluation.csv")